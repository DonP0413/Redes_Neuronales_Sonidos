from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = Path(
    os.getenv(
        "REPORT_OUTPUT",
        str(ROOT / "docs" / "INFORME_FINAL_RECONOCIMIENTO_AUDIO_ESC50.docx"),
    )
).resolve()

NAVY = "17365D"
BLUE = "2E74B5"
CYAN = "00A6B2"
GREEN = "2F855A"
DARK = "1F2937"
MUTED = "4B5563"
LIGHT = "F2F4F7"
PALE_BLUE = "EAF2F8"
PALE_GREEN = "EAF6EF"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != TABLE_WIDTH_DXA:
        raise ValueError(f"Los anchos deben sumar {TABLE_WIDTH_DXA}: {widths_dxa}")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "nil")


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Title", 30, NAVY, 0, 12),
        ("Subtitle", 14, MUTED, 0, 8),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, NAVY, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor.from_string(DARK)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)

    if "Caption" not in styles:
        styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True


def add_field(run, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, text, end):
        run._r.append(element)


def build_running_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(table, [6500, 2860])
    remove_table_borders(table)
    for cell in table.rows[0].cells:
        set_cell_shading(cell, WHITE)
        set_cell_margins(cell, top=20, bottom=30)
    left = table.cell(0, 0).paragraphs[0]
    left.paragraph_format.space_after = Pt(0)
    run = left.add_run("ONDA")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(NAVY)
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.paragraph_format.space_after = Pt(0)
    run = right.add_run("INFORME TÉCNICO FINAL")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    footer = section.footer
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(table, [7200, 2160])
    remove_table_borders(table)
    for cell in table.rows[0].cells:
        set_cell_shading(cell, WHITE)
        set_cell_margins(cell, top=30, bottom=20)
    left = table.cell(0, 0).paragraphs[0]
    left.paragraph_format.space_after = Pt(0)
    run = left.add_run("Sistemas Embebidos y Redes Neuronales · 2026_2")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.paragraph_format.space_after = Pt(0)
    run = right.add_run("Pág. ")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    page_run = right.add_run()
    add_field(page_run, "PAGE")

    first_footer = section.first_page_footer
    p = first_footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Proyecto académico · Agosto de 2026")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)


def add_paragraph(doc: Document, text: str, *, bold_lead: str | None = None, align=None, keep=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if bold_lead and text.startswith(bold_lead):
        first, rest = text[: len(bold_lead)], text[len(bold_lead) :]
        p.add_run(first).bold = True
        p.add_run(rest)
    else:
        p.add_run(text)
    p.paragraph_format.keep_together = keep
    return p


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def create_restarting_numbering(doc: Document) -> int:
    """Create a List Number instance that restarts at 1 for one logical list."""
    numbering = doc.part.numbering_part.element
    nums = numbering.findall(qn("w:num"))
    next_num_id = max(int(node.get(qn("w:numId"))) for node in nums) + 1
    base = next(node for node in nums if node.get(qn("w:numId")) == "5")
    abstract_id = base.find(qn("w:abstractNumId")).get(qn("w:val"))

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(next_num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_id)
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return next_num_id


def add_numbered(doc: Document, items: list[str]) -> None:
    num_id = create_restarting_numbering(doc)
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p_pr = p._p.get_or_add_pPr()
        old_num_pr = p_pr.find(qn("w:numPr"))
        if old_num_pr is not None:
            p_pr.remove(old_num_pr)
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        explicit_num_id = OxmlElement("w:numId")
        explicit_num_id.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl)
        num_pr.append(explicit_num_id)
        p_pr.append(num_pr)
        p.add_run(item)


def add_callout(doc: Document, title: str, text: str, *, fill=PALE_BLUE, accent=BLUE) -> None:
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [180, 9180])
    set_cell_shading(table.cell(0, 0), accent)
    set_cell_shading(table.cell(0, 1), fill)
    cell = table.cell(0, 1)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)
    p = cell.add_paragraph(text)
    p.paragraph_format.space_after = Pt(0)
    p.runs[0].font.color.rgb = RGBColor.from_string(DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int], *, first_col_bold=False):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, text in enumerate(headers):
        cell = header.cells[i]
        set_cell_shading(cell, LIGHT)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(NAVY)
        run.font.size = Pt(9)
    for row_values in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row_values):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            run.font.size = Pt(9)
            if first_col_bold and i == 0:
                run.bold = True
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_image(doc: Document, path: Path, caption: str, width_inches: float) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(str(path), width=Inches(width_inches))
    cap = doc.add_paragraph(caption, style="Caption")
    cap.paragraph_format.keep_with_next = False


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Inches(0.85)
    p.paragraph_format.space_after = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("INFORME FINAL")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(CYAN)

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Sistema web de reconocimiento de audio con redes neuronales")

    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Clasificación de seis sonidos ambientales reales mediante una CNN entrenada con ESC-50")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(28)
    run = p.add_run("CLAXON  ·  LADRIDO  ·  APLAUSOS  ·  LLUVIA  ·  SIRENA  ·  HELICÓPTERO")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2700, 6660])
    metadata = [
        ("Materia", "Sistemas Embebidos y Redes Neuronales"),
        ("Período", "2026_2"),
        ("Docente", "Ing. Sergio Granizo, MSc."),
        ("Integrantes", "Erick · Miguel · Jorge"),
        ("Defensa", "3 y 4 de septiembre de 2026"),
    ]
    for row, (label, value) in zip(table.rows, metadata):
        set_cell_shading(row.cells[0], LIGHT)
        p = row.cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(label.upper())
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(NAVY)
        p = row.cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(value)
        run.font.size = Pt(10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    run = p.add_run("ONDA")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor.from_string(NAVY)
    doc.add_page_break()


def add_contents(doc: Document) -> None:
    doc.add_heading("Resumen ejecutivo", level=1)
    add_paragraph(
        doc,
        "Este informe documenta el diseño, entrenamiento, evaluación e integración de un sistema web capaz de clasificar seis sonidos ambientales reales: claxon, ladrido, aplausos, lluvia, sirena y helicóptero. Se utilizaron 240 grabaciones de ESC-50 —40 por clase— y se respetaron sus folds oficiales para evitar fuga de información. La entrada de audio se transforma en un Mel-espectrograma de 64 × 320 × 1 y se procesa con una CNN ligera de 89.286 parámetros."
    )
    add_paragraph(
        doc,
        "Se compararon Adam y SGD con Momentum sobre la misma arquitectura, datos, semilla y presupuesto de 70 épocas. SGD fue seleccionado por lograr 100 % de accuracy de validación y 89,58 % de accuracy en prueba (43 aciertos de 48), con F1 macro de 89,29 %. El modelo resultante se integró en Flask con una interfaz para escuchar ejemplos, cargar o grabar audio y visualizar la clase junto con sus seis probabilidades."
    )
    add_callout(
        doc,
        "Resultado principal",
        "El 60 % del dataset se destinó al entrenamiento, 20 % a validación y 20 % a prueba. El conjunto de prueba permaneció aislado hasta la evaluación final.",
        fill=PALE_GREEN,
        accent=GREEN,
    )
    doc.add_heading("Contenido del informe", level=2)
    add_numbered(
        doc,
        [
            "Introducción y planteamiento del problema",
            "Objetivos general y específicos",
            "Desarrollo y arquitectura del sistema",
            "Dataset ESC-50, partición y preprocesamiento",
            "Arquitectura CNN y proceso de entrenamiento",
            "Resultados, conclusiones y recomendaciones",
        ],
    )
    doc.add_page_break()


def add_report_body(doc: Document) -> None:
    doc.add_heading("1. Introducción", level=1)
    add_paragraph(
        doc,
        "El reconocimiento automático de sonidos permite convertir una señal acústica en información útil para monitoreo, seguridad, accesibilidad e interacción humano–máquina. A diferencia del reconocimiento de voz, este proyecto se concentra en eventos ambientales que pueden ser identificados de forma intuitiva por el público antes de probar el sistema."
    )
    add_paragraph(
        doc,
        "El proyecto propone una solución completa: preparación de datos reales, transformación de audio, entrenamiento de una red neuronal convolucional, evaluación controlada e integración web. La aplicación denominada ONDA permite comparar la percepción humana con la predicción de la CNN y hace visible la distribución de probabilidades, no solo la etiqueta ganadora."
    )

    doc.add_heading("1.1 Planteamiento del problema", level=2)
    add_paragraph(
        doc,
        "Una señal de audio es una secuencia temporal extensa y sensible al volumen, al ruido y a la duración. Para clasificarla de forma consistente es necesario normalizarla, representar su contenido espectral y entrenar un modelo con ejemplos suficientemente variados. Además, una demostración académica debe ser reproducible y separar estrictamente los datos usados para entrenar de los usados para evaluar."
    )
    add_callout(
        doc,
        "Pregunta de trabajo",
        "¿Puede una CNN compacta distinguir seis eventos ambientales reales a partir de Mel-espectrogramas y ofrecer sus resultados en una aplicación web interactiva?",
    )

    doc.add_heading("1.2 Alcance", level=2)
    add_bullets(
        doc,
        [
            "Clasificación cerrada de seis clases: claxon, ladrido, aplausos, lluvia, sirena y helicóptero.",
            "Entrenamiento y evaluación con grabaciones reales de ESC-50.",
            "Comparación controlada entre Adam y SGD con Momentum.",
            "Aplicación web Flask con carga de archivos, micrófono, probabilidades y Mel-espectrograma.",
            "Preparación para despliegue académico en Railway mediante Docker y Gunicorn.",
        ],
    )

    doc.add_heading("2. Objetivos", level=1)
    doc.add_heading("2.1 Objetivo general", level=2)
    add_paragraph(
        doc,
        "Desarrollar e integrar un sistema web de reconocimiento de audio basado en una red neuronal convolucional que clasifique seis sonidos ambientales reales y presente al usuario la predicción, la confianza y la probabilidad de cada clase."
    )
    doc.add_heading("2.2 Objetivos específicos", level=2)
    add_bullets(
        doc,
        [
            "Construir un subconjunto balanceado y auditable de ESC-50 con 240 grabaciones reales.",
            "Aplicar un preprocesamiento idéntico durante entrenamiento e inferencia.",
            "Diseñar una CNN ligera para entradas Mel-espectrograma de 64 × 320 × 1.",
            "Comparar Adam y SGD con Momentum manteniendo constantes la arquitectura, los datos y la semilla.",
            "Medir accuracy, precision, recall, F1 macro y matriz de confusión sobre datos no vistos.",
            "Integrar el modelo seleccionado en una interfaz web clara y desplegable.",
        ],
    )
    doc.add_page_break()

    doc.add_heading("3. Desarrollo del sistema", level=1)
    add_paragraph(
        doc,
        "El desarrollo se organizó como un pipeline reproducible. Los datos se preparan una sola vez; el entrenamiento genera el modelo y las evidencias; el servidor reutiliza la misma función de preprocesamiento y el frontend consume una API JSON. Esta separación reduce inconsistencias entre laboratorio y producción."
    )
    doc.add_heading("3.1 Flujo funcional", level=2)
    add_numbered(
        doc,
        [
            "El usuario escucha ejemplos de las seis clases o aporta un audio propio.",
            "El backend valida el formato y limita el tamaño del archivo.",
            "La señal se convierte a mono, 16 kHz y cinco segundos.",
            "Se calcula y normaliza el Mel-espectrograma.",
            "La CNN produce seis probabilidades mediante Softmax.",
            "La interfaz muestra la clase principal, confianza, probabilidades y representación espectral.",
        ],
    )
    add_image(doc, ROOT / "results" / "architecture.png", "Figura 1. Arquitectura del modelo y flujo de tensores.", 6.25)

    doc.add_heading("3.2 Componentes implementados", level=2)
    add_table(
        doc,
        ["Componente", "Responsabilidad", "Tecnología"],
        [
            ["Preparación", "Selecciona clases, conserva folds y genera el manifiesto.", "Python"],
            ["Señal", "Carga, remuestrea, recorta, normaliza y calcula Mels.", "librosa / NumPy"],
            ["Modelo", "Entrena, compara optimizadores y exporta evidencias.", "TensorFlow / Keras"],
            ["API", "Valida audio, ejecuta inferencia y devuelve JSON.", "Flask"],
            ["Frontend", "Captura o carga audio y presenta resultados.", "HTML / CSS / JavaScript"],
            ["Despliegue", "Ejecuta un proceso web reproducible.", "Docker / Gunicorn / Railway"],
        ],
        [1800, 5040, 2520],
        first_col_bold=True,
    )

    doc.add_heading("4. Dataset ESC-50", level=1)
    add_paragraph(
        doc,
        "ESC-50 es un conjunto público de sonidos ambientales organizado en 50 categorías, con grabaciones de cinco segundos y cinco folds oficiales. Para este proyecto se escogieron seis categorías auditivamente reconocibles y se conservaron las 40 muestras disponibles de cada una. El subconjunto final contiene 240 WAV reales y balanceados."
    )
    add_table(
        doc,
        ["Clase de la aplicación", "Categoría ESC-50", "Muestras"],
        [
            ["Claxon", "car_horn", "40"],
            ["Ladrido", "dog", "40"],
            ["Aplausos", "clapping", "40"],
            ["Lluvia", "rain", "40"],
            ["Sirena", "siren", "40"],
            ["Helicóptero", "helicopter", "40"],
        ],
        [3600, 3840, 1920],
        first_col_bold=True,
    )
    doc.add_heading("4.1 Partición experimental", level=2)
    add_table(
        doc,
        ["Uso", "Folds", "Audios", "Porcentaje", "Por clase"],
        [
            ["Entrenamiento", "1–3", "144", "60 %", "24"],
            ["Validación", "4", "48", "20 %", "8"],
            ["Prueba", "5", "48", "20 %", "8"],
        ],
        [2550, 1350, 1620, 1920, 1920],
        first_col_bold=True,
    )
    add_callout(
        doc,
        "Control contra fuga de datos",
        "La selección del modelo se realizó únicamente con validación. Los 48 audios del fold 5 no intervinieron en el ajuste ni en la comparación de optimizadores y se usaron una sola vez para la evaluación final.",
        fill=PALE_GREEN,
        accent=GREEN,
    )
    add_paragraph(
        doc,
        "El manifiesto dataset/esc50/manifest.csv conserva la etiqueta, el fold, el archivo original y la referencia de Freesound. results/dataset_split.csv registra la asignación efectiva de cada muestra, lo que permite auditar el experimento. ESC-50 se distribuye bajo licencia CC BY-NC; por ello el uso y la demostración deben mantenerse académicos y no comerciales."
    )
    doc.add_page_break()

    doc.add_heading("5. Preprocesamiento y representación", level=1)
    add_paragraph(
        doc,
        "El audio no se entrega directamente a la red. Se transforma en una imagen tiempo–frecuencia que conserva patrones relevantes, como energía continua, impulsos breves y armónicos. El procedimiento es idéntico en entrenamiento y producción."
    )
    add_table(
        doc,
        ["Etapa", "Configuración", "Propósito"],
        [
            ["Canales", "Conversión a mono", "Unificar la forma de entrada."],
            ["Muestreo", "16.000 Hz", "Reducir costo y conservar el rango útil."],
            ["Duración", "5,0 s", "Recortar o completar con ceros."],
            ["Limpieza", "Recorte lateral + normalización por pico", "Disminuir silencios y diferencias de volumen."],
            ["STFT", "FFT 512; hop 256", "Obtener evolución espectral."],
            ["Mel", "64 bandas; 50–7.600 Hz", "Aproximar resolución auditiva."],
            ["Escala", "80 dB, normalizada a [0,1]", "Estabilizar el entrenamiento."],
            ["Tensor", "64 × 320 × 1", "Entrada fija para la CNN."],
        ],
        [1800, 2880, 4680],
        first_col_bold=True,
    )
    doc.add_heading("5.1 Aumento de datos", level=2)
    add_paragraph(
        doc,
        "Durante el entrenamiento se aplicaron variaciones leves de posición temporal, contraste y ruido gaussiano. Estas transformaciones se calculan al vuelo y no alteran los WAV originales. Se desactivan automáticamente durante validación, prueba e inferencia para que las métricas sean comparables."
    )

    doc.add_heading("6. Arquitectura CNN utilizada", level=1)
    add_paragraph(
        doc,
        "La red tiene 89.286 parámetros y combina extracción local de patrones con dos resúmenes temporales. El promedio temporal representa texturas sostenidas como lluvia o helicóptero; el máximo temporal preserva eventos cortos e intensos como claxon o aplausos. Ambos vectores se concatenan antes de la clasificación."
    )
    add_table(
        doc,
        ["Bloque", "Operación principal", "Salida conceptual"],
        [
            ["Entrada", "Mel-espectrograma normalizado", "64 × 320 × 1"],
            ["Convolucional 1", "Conv2D 16 + ReLU + MaxPool", "Patrones locales básicos"],
            ["Convolucional 2", "Conv2D 32 + ReLU + MaxPool", "Texturas intermedias"],
            ["Convolucional 3", "Conv2D 64 + ReLU + MaxPool", "Rasgos de alto nivel"],
            ["Resumen temporal", "Promedio + máximo en dos ramas", "Contexto sostenido y transitorio"],
            ["Cabeza", "Concatenación + Dense 64 + Dropout 0,35", "Vector discriminativo"],
            ["Salida", "Dense 6 + Softmax", "Probabilidad por clase"],
        ],
        [1950, 3900, 3510],
        first_col_bold=True,
    )
    add_callout(
        doc,
        "Por qué una CNN",
        "Las convoluciones comparten filtros sobre el espectrograma, detectan patrones sin depender de una posición exacta y requieren menos parámetros que una red totalmente conectada sobre todos los píxeles.",
    )
    doc.add_heading("6.1 Funciones de activación", level=2)
    add_paragraph(
        doc,
        "Las tres capas convolucionales y la capa Dense de 64 unidades utilizan ReLU. Esta activación conserva los valores positivos, anula los negativos y permite aprender relaciones no lineales con un cálculo sencillo. Frente a Sigmoid o Tanh, ReLU reduce la saturación en la zona positiva y disminuye el riesgo de gradientes muy pequeños en una red de esta profundidad."
    )
    add_paragraph(
        doc,
        "La capa final utiliza Softmax porque las seis categorías son mutuamente excluyentes. Softmax convierte los seis logits en valores entre 0 y 1 cuya suma es 1; la clase con mayor valor se presenta como predicción principal. No se utilizó Sigmoid en la salida porque no se trata de una clasificación multietiqueta."
    )
    doc.add_heading("6.2 Forward pass y backpropagation", level=2)
    add_paragraph(
        doc,
        "En el forward pass, cada lote de Mel-espectrogramas atraviesa los bloques convolucionales, las ramas de resumen temporal, la capa Dense y Softmax hasta producir seis probabilidades. La pérdida sparse categorical crossentropy compara esas probabilidades con la etiqueta real y cuantifica el error del lote."
    )
    add_paragraph(
        doc,
        "En la backpropagation, TensorFlow aplica diferenciación automática para calcular el gradiente de la pérdida con respecto a cada peso. Adam o SGD + Momentum utilizan esos gradientes para actualizar los parámetros en sentido contrario al error. Este ciclo se repite por cada batch y época; durante validación, prueba e inferencia solo se ejecuta el forward pass, sin modificar pesos."
    )
    doc.add_page_break()

    doc.add_heading("7. Entrenamiento y diseño experimental", level=1)
    add_paragraph(
        doc,
        "Se entrenaron dos copias de la misma CNN para aislar el efecto del optimizador. La arquitectura, los folds, el preprocesamiento, el batch, la semilla y el máximo de épocas permanecieron constantes. La función de pérdida fue entropía cruzada categórica dispersa y la métrica de seguimiento fue accuracy."
    )
    add_table(
        doc,
        ["Parámetro", "Valor"],
        [
            ["Semilla", "2026"],
            ["Batch", "16"],
            ["Máximo de épocas", "70 por experimento"],
            ["Pérdida", "Sparse categorical crossentropy"],
            ["Adam", "Learning rate inicial 0,001"],
            ["SGD + Momentum", "Learning rate inicial 0,01; momentum 0,9"],
            ["Early stopping", "Paciencia 12; restauración de los mejores pesos"],
            ["Reducción de LR", "ReduceLROnPlateau; factor 0,5"],
            ["Regla de selección", "Mayor accuracy de validación; luego menor loss"],
        ],
        [3060, 6300],
        first_col_bold=True,
    )
    doc.add_heading("7.1 Épocas ejecutadas", level=2)
    add_paragraph(
        doc,
        "Adam y SGD completaron las 70 épocas configuradas porque sus mejores resultados aparecieron cerca del final del presupuesto. Adam alcanzó su mejor validación en la época 60; SGD + Momentum, en la época 62. Para evaluar se restauraron los pesos de esas mejores épocas, no necesariamente los de la época 70."
    )
    doc.add_heading("7.2 Comparación de optimizadores", level=2)
    add_image(doc, ROOT / "results" / "optimizer_comparison.png", "Figura 2. Curvas de accuracy y loss para Adam y SGD + Momentum.", 6.35)
    add_paragraph(
        doc,
        "La decisión se tomó antes de consultar las etiquetas del test. SGD logró 100 % de accuracy de validación y una loss de 0,0883, frente a 97,92 % y 0,0992 de Adam. Por la regla definida, SGD fue promovido como modelo final."
    )
    doc.add_heading("7.3 Diagnóstico del aprendizaje", level=2)
    add_bullets(
        doc,
        [
            "Underfitting: no fue el problema dominante. Las curvas de entrenamiento y validación alcanzaron accuracies altas y la loss descendió de forma sostenida, por lo que la CNN sí tuvo capacidad para aprender el patrón de las seis clases.",
            "Overfitting: existe un riesgo moderado por disponer de solo 40 audios por clase y por la diferencia entre 100 % de validación y 89,58 % de prueba. Se mitigó con aumento de datos, dropout 0,35, reducción del learning rate, early stopping y partición por folds.",
            "Vanishing gradient: no se observaron señales claras en las curvas, aunque no se midieron los gradientes de forma directa. La red es relativamente corta, usa ReLU y entradas normalizadas; además, la mejora sostenida de accuracy y loss indica que el gradiente llegó a las capas iniciales durante el entrenamiento.",
        ],
    )

    doc.add_heading("8. Resultados y análisis", level=1)
    add_table(
        doc,
        ["Optimizador", "Mejor época", "Accuracy val.", "Loss val.", "Accuracy test", "F1 macro"],
        [
            ["Adam", "60", "97,92 %", "0,0992", "87,50 %", "87,17 %"],
            ["SGD + Momentum", "62", "100 %", "0,0883", "89,58 %", "89,29 %"],
        ],
        [2040, 1320, 1560, 1320, 1680, 1440],
        first_col_bold=True,
    )
    add_callout(
        doc,
        "Modelo seleccionado: SGD + Momentum",
        "43 aciertos de 48 audios de prueba · Accuracy 89,58 % · Precision macro 92,12 % · Recall macro 89,58 % · F1 macro 89,29 %.",
        fill=PALE_GREEN,
        accent=GREEN,
    )
    doc.add_heading("8.1 Rendimiento por clase", level=2)
    add_table(
        doc,
        ["Clase", "Precision", "Recall", "F1", "Aciertos / 8"],
        [
            ["Claxon", "72,73 %", "100 %", "84,21 %", "8"],
            ["Ladrido", "100 %", "75 %", "85,71 %", "6"],
            ["Aplausos", "100 %", "62,50 %", "76,92 %", "5"],
            ["Lluvia", "80 %", "100 %", "88,89 %", "8"],
            ["Sirena", "100 %", "100 %", "100 %", "8"],
            ["Helicóptero", "100 %", "100 %", "100 %", "8"],
        ],
        [2500, 1700, 1600, 1600, 1960],
        first_col_bold=True,
    )
    doc.add_page_break()

    doc.add_heading("8.2 Matriz de confusión", level=2)
    add_image(doc, ROOT / "results" / "confusion_matrix.png", "Figura 3. Matriz de confusión del modelo SGD sobre 48 audios de prueba.", 5.55)
    add_paragraph(
        doc,
        "Sirena y helicóptero fueron reconocidos sin errores. Claxon y lluvia alcanzaron recall de 100 %, aunque recibieron falsos positivos y por ello su precision fue menor. Aplausos presentó el recall más bajo: tres audios se confundieron, principalmente con texturas continuas. Esto sugiere que el modelo identifica bien eventos acústicamente distintivos, pero necesita mayor variedad para las clases con patrones irregulares."
    )
    add_paragraph(
        doc,
        "La accuracy de validación superior a la de prueba no debe interpretarse como falla del sistema. Cada subconjunto contiene grabaciones distintas y el test puede incluir condiciones más difíciles. La métrica que debe comunicarse como rendimiento final es 89,58 % en prueba, no el 100 % de validación."
    )

    doc.add_heading("9. Aplicación web e inferencia", level=1)
    add_paragraph(
        doc,
        "El modelo se carga una sola vez al iniciar Flask. El usuario puede seleccionar un archivo o grabar desde el micrófono; el backend valida el tamaño y la extensión, guarda un temporal, aplica el preprocesamiento compartido y elimina el archivo al finalizar. La respuesta incluye etiqueta, confianza, seis probabilidades, tiempo aproximado y datos del espectrograma."
    )
    add_table(
        doc,
        ["Elemento", "Comportamiento"],
        [
            ["Ejemplos", "Seis audios reales del fold 5 permiten conocer las clases antes de probar."],
            ["Entrada", "Carga WAV/FLAC/MP3/OGG/M4A/WEBM o grabación de micrófono."],
            ["Resultado", "Clase principal, anillo de confianza y barras para las seis probabilidades."],
            ["Evidencia visual", "Forma de onda y Mel-espectrograma del audio analizado."],
            ["API", "POST /api/predict; GET /api/status; GET /health."],
            ["Producción", "Docker, Python 3.11 y Gunicorn con un worker para controlar memoria."],
        ],
        [2400, 6960],
        first_col_bold=True,
    )
    add_callout(
        doc,
        "Interpretación correcta de la confianza",
        "Softmax reparte 100 % entre las seis clases conocidas. Una confianza alta no demuestra que el sonido pertenezca realmente a una de ellas; la versión actual no incluye una clase “desconocido”.",
    )

    doc.add_heading("10. Limitaciones", level=1)
    add_bullets(
        doc,
        [
            "Solo existen 40 grabaciones por clase; la diversidad acústica es limitada.",
            "La evaluación utiliza un único fold de prueba de 48 audios.",
            "La clasificación es cerrada: todo audio recibe una de las seis etiquetas.",
            "Ruido intenso, distancia al micrófono o sonidos superpuestos pueden reducir el rendimiento.",
            "Las probabilidades no han sido calibradas como una medida estadística de certeza.",
            "La licencia CC BY-NC de ESC-50 restringe el uso a contextos no comerciales con atribución.",
        ],
    )
    doc.add_page_break()

    doc.add_heading("11. Conclusiones", level=1)
    add_numbered(
        doc,
        [
            "Se completó un sistema de extremo a extremo que une dataset real, procesamiento de señal, CNN, evaluación e interfaz web.",
            "La partición 60 % entrenamiento, 20 % validación y 20 % prueba, basada en folds oficiales, permitió evaluar sin reutilizar muestras entre etapas.",
            "La arquitectura de 89.286 parámetros fue suficiente para alcanzar 89,58 % de accuracy y 89,29 % de F1 macro sobre prueba.",
            "SGD + Momentum superó a Adam bajo el protocolo fijado y fue seleccionado en la época 62 por su mejor resultado de validación.",
            "La interfaz hace interpretable la salida al mostrar las seis probabilidades y permite que el público escuche las clases antes de realizar la demostración.",
            "El rendimiento perfecto de algunas clases y la dificultad en aplausos evidencian que la separabilidad acústica y la variedad de ejemplos influyen directamente en el resultado.",
        ],
    )

    doc.add_heading("12. Recomendaciones", level=1)
    add_bullets(
        doc,
        [
            "Ampliar el dataset con grabaciones propias y condiciones reales de Ecuador, manteniendo conjuntos separados por fuente.",
            "Ejecutar validación cruzada sobre los cinco folds para estimar variabilidad y no depender de una única partición.",
            "Agregar una clase desconocido o un umbral calibrado para rechazar audios fuera del dominio.",
            "Reforzar especialmente aplausos y ladrido mediante más ejemplos, mezcla controlada de ruido y pruebas de micrófono.",
            "Medir latencia, memoria y tamaño del modelo si se migra a un dispositivo embebido; evaluar cuantización TensorFlow Lite.",
            "Mantener versionados el modelo, metadata, manifiesto y métricas para que toda cifra de la defensa sea rastreable.",
            "En la exposición, comunicar 89,58 % como accuracy final de prueba y explicar que 100 % corresponde únicamente a validación.",
        ],
    )

    doc.add_heading("13. Reproducibilidad", level=1)
    add_paragraph(doc, "El experimento puede reconstruirse con los scripts incluidos en el repositorio:")
    add_table(
        doc,
        ["Acción", "Comando"],
        [
            ["Preparar ESC-50", "python -m src.prepare_esc50 --source-dir C:\\tmp\\ESC-50"],
            ["Entrenar", "python -m src.train --dataset-dir dataset\\esc50 --epochs 70 --batch-size 16 --patience 12 --seed 2026"],
            ["Probar", "python -m pytest -q"],
            ["Servidor local", "python app.py"],
        ],
        [2400, 6960],
        first_col_bold=True,
    )
    add_paragraph(
        doc,
        "Los artefactos clave son model/audio_classifier.keras, model/metadata.json, results/optimizer_metrics.csv, results/classification_report.json, results/test_predictions.csv y las figuras incluidas en este informe."
    )

    doc.add_heading("Referencias", level=1)
    add_numbered(
        doc,
        [
            "Piczak, K. J. (2015). ESC: Dataset for Environmental Sound Classification. Proceedings of ACM Multimedia. Repositorio ESC-50: https://github.com/karolpiczak/ESC-50",
            "TensorFlow / Keras. Documentación oficial: https://www.tensorflow.org/",
            "librosa. Audio and music signal analysis in Python: https://librosa.org/",
            "Flask. Documentación oficial: https://flask.palletsprojects.com/",
        ],
    )
    add_paragraph(
        doc,
        "Fuente de resultados: artefactos generados por el entrenamiento del proyecto el 23 de agosto de 2026. Las cifras se expresan con redondeo a dos decimales; los valores completos permanecen en los archivos JSON y CSV del repositorio.",
    )


def build_document() -> None:
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    build_running_header_footer(doc)
    props = doc.core_properties
    props.title = "Informe final — Sistema web de reconocimiento de audio con redes neuronales"
    props.subject = "Clasificación de sonidos ambientales ESC-50 con CNN"
    props.author = "Erick, Miguel y Jorge"
    props.keywords = "ESC-50, CNN, audio, Mel-espectrograma, Flask, Railway"
    props.comments = "Informe generado a partir de los artefactos reales del proyecto ONDA."
    add_cover(doc)
    add_contents(doc)
    add_report_body(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
